import docx

d = docx.Document('/Users/ionbota/Downloads/demo.docx')
print(d.paragraphs) # all paragraph objects
print(d.paragraphs[0].text) # Show first paragraph
print(d.paragraphs[1].text) # Show second paragraph

p = d.paragraphs[1]
print(p.runs[0].text)  # show how many runs are, runs are considered when style is changed
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

print(p.runs[1].bold)   # Each of the variables has Bold, Italic, Underline valiables
print(p.runs[3].italic)

# Change style text and save to new word
p.runs[3].text = 'italic and underline.'
d.save('/Users/ionbota/Downloads/demo2.docx')

# Change paragraph syle
p.style = 'Title'
d.save('/Users/ionbota/Downloads/demo3.docx')

# Create new word
d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('Another paragraph')
d.save('/Users/ionbota/Downloads/demo4.docx')

# Add more text
p = d.paragraphs[0]
p.add_run('This is new run')
p.runs
p.runs[1].bold = True
p.runs[1].italic = True
p.runs[1].underline = True
d.save('/Users/ionbota/Downloads/demo5.docx')

#---------#----------#---------#----------#

# Open all text in document
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('/Users/ionbota/Downloads/demo.docx'))
